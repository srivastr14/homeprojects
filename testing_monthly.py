import pandas as pd
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE


# Download from PuSH (Monthly Email)
# Data>Filter for first row
# Select "." frmo Proposals column in filter and then clear contents
# Select all instru except blanks and paint cell then clear filters in dropdown
# Sort JIF descending, sort instr by cell color
# Clear formats
# Save as monthly_jif.csv

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

df = pd.read_csv('monthly_jif.csv')
df = df.fillna('')
df.rename(columns={'Citation_PuSH_style___________________________________________': 'Citation'}, inplace=True)

doc = Document()
z = 0

for y, x in df.iterrows():
    z += 1
    print(z)
    p = doc.add_paragraph()
    if (x[2] == '' or int(x[2]) < 7) and x[3] == '':
        citation = p.add_run(str(x[0]))
        citation.add_break()
    else:
        p.add_run(str(x[0]) + " ")
    if x[2] and int(x[2]) > 7:
        jif_raw = "JIF: " + str(x[2])
        jif = p.add_run(jif_raw)
        jif.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        jif.bold = True
    if x[3] != "":
        if ',' in str(x[3]):
            no_space = str(x[3])
            space = no_space.replace(',',', ')
            linked_raw = ' Linked: ' + space
        else:
            linked_raw = ' Linked: ' + x[3]
        linked = p.add_run(linked_raw)
        linked.font.color.rgb = RGBColor(255,0,0)
        linked.add_break()
    elif x[3] == '':
        jif.add_break()
    if x[4] != '':
        instr_raw = str(x[4])
        instr = p.add_run(instr_raw)
        instr.bold = True
    if x[4] != '' and x[5] != '':
        noninstr_raw = str(x[5])
        noninstr = p.add_run(' ' + noninstr_raw)
        noninstr.bold = True
    if x[5] != '' and x[4] == '':
        noninstr_raw = str(x[5])
        noninstr = p.add_run(noninstr_raw)
        noninstr.bold = True
    if x[6]:
        # doi = p.add_run(" " + str(x[6]))
        p.add_run(' ')
        add_hyperlink(p, str(x[6]), str(x[6]))
        p.add_run().add_break()

doc.save('monthlyPuSH.docx')

